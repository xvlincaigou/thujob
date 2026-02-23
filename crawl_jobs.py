#!/usr/bin/env python3
"""
清华大学就业信息网爬虫
支持按日期范围爬取招聘信息，并获取详情页完整内容
"""

import requests
from bs4 import BeautifulSoup
from datetime import datetime, date
import json
import re
import time
import sys
import os
from urllib.parse import urljoin

# 可选的 docx 导出
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


BASE_URL = "https://career.cic.tsinghua.edu.cn/xsglxt/f/jyxt/anony/xxfb"
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Content-Type": "application/x-www-form-urlencoded",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
}


def parse_job_item(li_element):
    """解析单个招聘信息（列表页）"""
    job = {}

    span = li_element.find('span')
    if span:
        job['date'] = span.get_text(strip=True)

    a_tag = li_element.find('a')
    if a_tag:
        href = a_tag.get('ahref', '')
        if href:
            job['url'] = urljoin("https://career.cic.tsinghua.edu.cn", href)

        full_text = a_tag.get_text(strip=True)
        job['full_text'] = full_text

        if '————' in full_text:
            parts = full_text.split('————', 1)
            job['title'] = parts[0].strip()
            job['company'] = parts[1].strip()
        else:
            job['title'] = full_text
            job['company'] = ''

        job['scope'] = a_tag.get('fbfw', '')
        style = a_tag.get('style', '')
        job['is_highlighted'] = 'ff0000' in style

    return job


def fetch_detail_page(url, delay=1):
    """获取详情页内容"""
    try:
        time.sleep(delay)
        response = requests.get(url, headers=HEADERS, timeout=30)
        response.encoding = 'utf-8'
        return response.text
    except Exception as e:
        print(f"    获取详情页失败: {e}")
        return None


def parse_detail_page(html):
    """解析详情页，提取完整内容文本"""
    soup = BeautifulSoup(html, 'html.parser')

    # 查找主要内容区域
    content_div = soup.find('div', class_='content teacher') or soup.find('div', class_='content')

    if not content_div:
        return {}

    # 返回完整文本内容
    return {
        'full_content': content_div.get_text('\n', strip=True)
    }


def save_to_docx(jobs, filename):
    """保存招聘信息到 Word 文档"""
    doc = Document()

    # 设置文档标题
    title = doc.add_heading('清华大学就业信息网招聘信息', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加统计信息
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'共 {len(jobs)} 条招聘信息')
    doc.add_paragraph()

    for i, job in enumerate(jobs, 1):
        # 添加分隔线（除第一条外）
        if i > 1:
            doc.add_paragraph('_' * 60)

        # 添加标题（职位名 - 公司）
        heading_text = f"{i}. {job.get('title', '未知职位')}"
        heading = doc.add_heading(heading_text, level=2)

        # 基本信息表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        # 添加基本信息行
        basic_info = [
            ('发布日期', job.get('date', 'N/A')),
            ('公司名称', job.get('company', 'N/A')),
            ('发布范围', job.get('scope', 'N/A')),
            ('详情链接', job.get('url', 'N/A')),
        ]

        for label, value in basic_info:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        # 如果有详情内容，添加详情
        if job.get('detail') and job['detail'].get('full_content'):
            doc.add_paragraph()
            detail_heading = doc.add_heading('详细信息', level=3)
            detail_para = doc.add_paragraph(job['detail']['full_content'])

        # 添加空行
        doc.add_paragraph()

    # 保存文档
    doc.save(filename)


def fetch_page(page_no):
    """获取指定页的数据"""
    data = {
        "pgno": page_no,
        "type": "",
        "flag": "",
        "zwmc": "",
        "dwmc": "",
        "gzdqmc": "",
    }

    try:
        response = requests.post(BASE_URL, data=data, headers=HEADERS, timeout=30)
        response.encoding = 'utf-8'
        return response.text
    except Exception as e:
        print(f"获取第 {page_no} 页失败: {e}")
        return None


def parse_page(html):
    """解析页面HTML，提取招聘信息和分页信息"""
    soup = BeautifulSoup(html, 'html.parser')
    jobs = []

    list_items = soup.find_all('li', class_='clearfix')

    for li in list_items:
        span = li.find('span')
        a_tag = li.find('a')
        if span and a_tag and re.match(r'\d{4}-\d{2}-\d{2}', span.get_text(strip=True)):
            job = parse_job_item(li)
            if job:
                jobs.append(job)

    pagination = {}
    total_pg_elem = soup.find('b', id='totalPg')
    page_no_elem = soup.find('b', id='pageNo')

    if total_pg_elem:
        pagination['total_pages'] = int(total_pg_elem.get_text(strip=True))
    if page_no_elem:
        pagination['current_page'] = int(page_no_elem.get_text(strip=True))

    return jobs, pagination


def parse_date(date_str):
    """将日期字符串解析为date对象"""
    try:
        return datetime.strptime(date_str, '%Y-%m-%d').date()
    except:
        return None


def crawl_jobs(start_date=None, end_date=None, max_pages=None, output_file=None, output_basename=None,
               delay=1, fetch_details=False, detail_delay=1, progress_interval=10):
    """
    爬取招聘信息

    Args:
        start_date: 开始日期
        end_date: 结束日期
        max_pages: 最大爬取页数
        output_file: 输出文件路径（命令行模式使用）
        output_basename: 输出文件基础名（交互模式使用，同时生成json和docx）
        delay: 列表页请求间隔秒数
        fetch_details: 是否获取详情页
        detail_delay: 详情页请求间隔秒数
        progress_interval: 进度报告间隔（每处理N条报告一次）
    """
    if isinstance(start_date, str):
        start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
    if isinstance(end_date, str):
        end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

    if start_date:
        print(f"开始日期: {start_date}")
    if end_date:
        print(f"结束日期: {end_date}")
    print(f"列表页请求间隔: {delay}秒")
    if fetch_details:
        print(f"详情页请求间隔: {detail_delay}秒")
        print("注意：获取详情页会显著增加爬取时间")
    print("-" * 50)

    all_jobs = []
    page_no = 1
    stop_crawling = False

    while not stop_crawling:
        if max_pages and page_no > max_pages:
            print(f"已达到最大页数限制: {max_pages}")
            break

        print(f"正在获取第 {page_no} 页...")

        html = fetch_page(page_no)
        if not html:
            break

        jobs, pagination = parse_page(html)
        total_pages = pagination.get('total_pages', 0)

        print(f"  获取到 {len(jobs)} 条招聘信息")

        out_of_range_count = 0

        for i, job in enumerate(jobs):
            job_date = parse_date(job.get('date', ''))

            if job_date:
                if end_date and job_date > end_date:
                    continue
                if start_date and job_date < start_date:
                    out_of_range_count += 1
                    continue

            # 获取详情页
            if fetch_details and job.get('url'):
                if (i + 1) % progress_interval == 0:
                    print(f"    正在获取第 {i+1}/{len(jobs)} 条详情...")

                detail_html = fetch_detail_page(job['url'], delay=detail_delay)
                if detail_html:
                    job['detail'] = parse_detail_page(detail_html)

            all_jobs.append(job)

        if start_date and out_of_range_count > 0 and out_of_range_count >= len(jobs) * 0.8:
            stop_crawling = True

        if page_no >= total_pages:
            print("已到达最后一页")
            break

        if stop_crawling:
            print("已超出日期范围，停止爬取")
            break

        page_no += 1
        time.sleep(delay)

    print("-" * 50)
    print(f"共爬取 {len(all_jobs)} 条符合要求的招聘信息")

    all_jobs.sort(key=lambda x: x.get('date', ''), reverse=True)

    if output_basename:
        # 交互模式：同时生成 json 和 docx
        json_file = f"{output_basename}.json"
        docx_file = f"{output_basename}.docx"

        # 保存为 JSON
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(all_jobs, f, ensure_ascii=False, indent=2)
        print(f"JSON 结果已保存到: {json_file}")

        # 保存为 DOCX
        if DOCX_AVAILABLE:
            save_to_docx(all_jobs, docx_file)
            print(f"Word 文档已保存到: {docx_file}")
        else:
            print("警告：python-docx 库未安装，跳过生成 Word 文档")

    elif output_file:
        # 命令行模式：根据文件扩展名决定输出格式
        if output_file.lower().endswith('.docx'):
            if DOCX_AVAILABLE:
                save_to_docx(all_jobs, output_file)
                print(f"结果已保存到: {output_file}")
            else:
                print("错误：python-docx 库未安装，无法生成 docx 文件")
                print("请运行: pip install python-docx")
                # 回退到 json 格式
                json_file = output_file.rsplit('.', 1)[0] + '.json'
                with open(json_file, 'w', encoding='utf-8') as f:
                    json.dump(all_jobs, f, ensure_ascii=False, indent=2)
                print(f"已改为保存为 JSON 格式: {json_file}")
        else:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(all_jobs, f, ensure_ascii=False, indent=2)
            print(f"结果已保存到: {output_file}")
    else:
        print("\n结果预览（前5条）：")
        for job in all_jobs[:5]:
            print(f"\n[{job.get('date', 'N/A')}] {job.get('title', 'N/A')}")
            print(f"  公司: {job.get('company', 'N/A')}")
            if job.get('detail'):
                detail = job['detail']
                if detail.get('full_content'):
                    print(f"  [有详细内容]")

    return all_jobs


def interactive_mode():
    """交互模式"""
    print("=" * 60)
    print("清华大学就业信息网招聘信息爬虫")
    print("=" * 60)
    print()

    print("本工具可以帮助你爬取清华大学就业信息网上的招聘信息。")
    print("网站地址: https://career.cic.tsinghua.edu.cn/xsglxt/f/jyxt/anony/xxfb")
    print()
    print("提示：")
    print("  - 网站上的日期通常是招聘活动的日期，可能是未来的日期")
    print("  - 如果未指定日期，默认爬取最近一个月的招聘信息")
    print("  - 获取详情页会显著增加爬取时间，请谨慎选择")
    print()

    # 输入开始日期
    while True:
        print("-" * 40)
        start_input = input("请输入开始日期 (格式: YYYY-MM-DD，直接回车表示今天): ").strip()
        if not start_input:
            start_date = date.today()
            break
        try:
            start_date = datetime.strptime(start_input, '%Y-%m-%d').date()
            break
        except ValueError:
            print("日期格式错误，请使用 YYYY-MM-DD 格式")

    # 输入结束日期
    while True:
        end_input = input("请输入结束日期 (格式: YYYY-MM-DD，直接回车表示今天): ").strip()
        if not end_input:
            end_date = date.today()
            break
        try:
            end_date = datetime.strptime(end_input, '%Y-%m-%d').date()
            break
        except ValueError:
            print("日期格式错误，请使用 YYYY-MM-DD 格式")

    # 是否获取详情页
    print()
    fetch_details_input = input("是否获取详情页内容？(y/N，默认否): ").strip().lower()
    fetch_details = fetch_details_input in ('y', 'yes', '是')

    # 最大页数
    print()
    max_pages_input = input("最大爬取页数 (直接回车表示不限制): ").strip()
    max_pages = int(max_pages_input) if max_pages_input.isdigit() else None

    # 固定请求间隔为0.5秒
    delay = 0.5
    detail_delay = 0.5

    # 输出文件基础名（不带扩展名）
    print()
    default_basename = f"jobs_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    basename_input = input(f"输出文件基础名 (默认: {default_basename}): ").strip()
    basename = basename_input if basename_input else default_basename

    print()
    print("=" * 60)
    print("配置确认：")
    print(f"  开始日期: {start_date}")
    print(f"  结束日期: {end_date}")
    print(f"  获取详情页: {'是' if fetch_details else '否'}")
    print(f"  最大页数: {max_pages or '不限制'}")
    print(f"  输出文件: {basename}.json 和 {basename}.docx")
    print("=" * 60)
    print()

    confirm = input("确认开始爬取？(Y/n): ").strip().lower()
    if confirm in ('n', 'no', '否'):
        print("已取消")
        return

    print()
    crawl_jobs(
        start_date=start_date,
        end_date=end_date,
        max_pages=max_pages,
        output_basename=basename,
        delay=delay,
        fetch_details=fetch_details,
        detail_delay=detail_delay
    )

    print()
    print(f"爬取完成！结果已保存到:")
    print(f"  - {basename}.json")
    print(f"  - {basename}.docx")


def main():
    import argparse

    parser = argparse.ArgumentParser(description='清华大学就业信息网招聘信息爬虫')
    parser.add_argument('--start-date', '-s', type=str, help='开始日期 (格式: YYYY-MM-DD)')
    parser.add_argument('--end-date', '-e', type=str, help='结束日期 (格式: YYYY-MM-DD)')
    parser.add_argument('--max-pages', '-m', type=int, help='最大爬取页数')
    parser.add_argument('--output', '-o', type=str, help='输出文件路径 (支持 .json 或 .docx 格式，根据扩展名自动判断)')
    parser.add_argument('--delay', '-d', type=float, default=1, help='请求间隔秒数 (默认: 1)')
    parser.add_argument('--fetch-details', '-f', action='store_true', help='获取详情页内容')
    parser.add_argument('--detail-delay', type=float, default=1, help='详情页请求间隔秒数 (默认: 1)')
    parser.add_argument('--cli', '-c', action='store_true', help='强制使用命令行模式（非交互式）')

    args = parser.parse_args()

    # 如果没有提供任何参数，进入交互模式
    if not args.cli and len(sys.argv) == 1:
        interactive_mode()
        return

    # 命令行模式
    if not args.start_date and not args.end_date:
        end_date = date.today()
        start_date = date(end_date.year, end_date.month, 1)
        print(f"未指定日期范围，默认爬取本月招聘信息 ({start_date} 至 {end_date})")
    else:
        start_date = args.start_date
        end_date = args.end_date

    crawl_jobs(
        start_date=start_date,
        end_date=end_date,
        max_pages=args.max_pages,
        output_file=args.output,
        delay=args.delay,
        fetch_details=args.fetch_details,
        detail_delay=args.detail_delay
    )


if __name__ == '__main__':
    main()
